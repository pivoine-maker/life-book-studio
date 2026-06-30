#!/usr/bin/env python3
from __future__ import annotations

import argparse
import copy
import json
import os
import re
import sys
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from openai import AzureOpenAI
from openpyxl import load_workbook
from docx import Document


DEFAULT_ENDPOINT = "https://aidp.bytedance.net/api/modelhub/online/v2/crawl"
DEFAULT_API_VERSION = "2024-02-01"
DEFAULT_MODEL = "gpt-5.4-2026-03-05"
DEFAULT_MAX_REPLACEMENT_LENGTH = 120
DEFAULT_MAX_EDIT_RATIO = 0.4
DEFAULT_MAX_SECTION_CHARS = 6000
URL_RE = re.compile(r"(?:https?://|www\.)\S+", re.IGNORECASE)
GARBLED_RE = re.compile(r"[�◆★☆◎◇□△※→←↘↗◉◌]{2,}")
BUTTON_WORDS = (
    "点击",
    "收藏",
    "关注",
    "上一章",
    "下一章",
    "返回顶部",
    "加入书架",
    "立即下载",
    "APP下载",
)
TRAILING_PATTERNS = [
    re.compile(r"备案号[:：]?\s*[A-Za-z0-9\-]+", re.IGNORECASE),
    re.compile(r"版权所有.*"),
    re.compile(r"本章未完.*"),
]
TITLEISH_RE = re.compile(r"^(作者[:：].*|书名[:：].*|简介[:：].*|人物小传[:：]?.*|人物介绍[:：]?.*|剧情大纲[:：]?.*|故事梗概[:：]?.*|主要人物[:：]?.*|人物设定[:：]?.*|角色介绍[:：]?.*|信息流[:：]?.*|看点[:：]?.*|定位[:：]?.*|一句话介绍[:：]?.*|项目[:：].*|片名[:：].*|来源[:：].*|题材[:：].*|集数[:：].*|时长[:：].*|状态[:：].*|报价[:：].*)$")
CHAPTER_MARK_RE = re.compile(r"^(第\s*[0-9一二三四五六七八九十百千两零]+\s*[集章节幕回卷]\b.*|分镜\s*[0-9一二三四五六七八九十百千两零]+.*|第\s*[0-9一二三四五六七八九十百千两零]+\s*场\b.*)$")
BODY_START_MARK_RE = re.compile(r"^(第\s*[0-9一二三四五六七八九十百千两零]+\s*[集章节幕回卷]\b.*|分镜\s*[0-9一二三四五六七八九十百千两零]+.*|第\s*[0-9一二三四五六七八九十百千两零]+\s*场\b.*|[0-9]{1,3}[.、]?$)")
METADATA_BLOCK_START_RE = re.compile(r"^(知乎盐选\s*\|.*|无障碍|查看详情|剧本大纲|大纲|人物小传|人物介绍|剧情大纲|故事梗概|主要人物|人物设定|角色介绍|信息流|项目|片名|影片名|剧名|来源|年代|题材|故事题材|电影类型|影片时长|字数|集数|时长|状态|报价|一句话介绍|故事看点|看点|定位|微视频剧本.*投稿|院线电影剧本.*)[:：]?")


@dataclass
class Section:
    section_id: str
    text: str
    kind: str
    locator: dict[str, Any]
    original_text: str | None = None
    preprocess_notes: list[str] | None = None


@dataclass
class MaterialDocument:
    source_path: Path
    source_type: str
    sections: list[Section]
    workbook: Any | None = None
    sheet_order: list[str] | None = None
    editable_columns: dict[str, list[int]] | None = None
    header_map: dict[str, dict[int, str]] | None = None
    selection: dict[str, Any] | None = None
    body_start_section_id: str | None = None


class CleaningError(Exception):
    pass


class AzureCleaningClient:
    def __init__(self, api_key: str, endpoint: str, api_version: str, model: str, logid: str | None):
        headers = {"X-TT-LOGID": logid or "material-cleaner"}
        self.client = AzureOpenAI(
            api_key=api_key,
            api_version=api_version,
            azure_endpoint=endpoint,
            default_headers=headers,
        )
        self.model = model

    def analyze_document(self, document: MaterialDocument) -> dict[str, Any]:
        section_map = {section.section_id: section for section in document.sections}
        overview_payload = {
            "source_type": document.source_type,
            "section_count": len(document.sections),
            "section_previews": [
                {
                    "section_id": section.section_id,
                    "kind": section.kind,
                    "locator": section.locator,
                    "preview": section.text[:200],
                }
                for section in document.sections
            ],
        }
        overview_prompt = (
            "你是小说/剧本素材清洗审校助手。你必须先理解整篇素材。"
            "目标是只保留正文，删除人物小传、人物介绍、剧情大纲、故事梗概、项目包装信息、导流信息等非正文内容。"
            "但必须保留章节号、集号、幕号、节号、分镜号，例如第1集、第3章、第2幕、第4节、分镜5。"
            "如果正文存在明确起点（如单独一行的1、01、001、1.、第1集、第一章、分镜1），必须优先定位正文起点，正文起点之前的全部文本都应删除。"
            "请识别需要处理的 section，并说明问题类别。"
            "只输出 JSON："
            '{"target_sections":[{"section_id":"...","categories":["leading_metadata"],"reason":"...","manual_review":false}]}'
        )
        overview_response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": "只输出 JSON，不要输出解释。"},
                {"role": "user", "content": [{"type": "text", "text": overview_prompt}]},
                {"role": "user", "content": [{"type": "text", "text": json.dumps(overview_payload, ensure_ascii=False)}]},
            ],
            max_tokens=2500,
            stream=False,
            temperature=0,
        )
        overview_data = parse_json_object(extract_response_text(overview_response.model_dump()))
        targets = overview_data.get("target_sections") or []
        operations: list[dict[str, Any]] = []
        issues: list[dict[str, Any]] = []
        deletion_summaries: list[dict[str, Any]] = []
        for target in targets:
            if not isinstance(target, dict):
                continue
            section_id = target.get("section_id")
            if not isinstance(section_id, str) or section_id not in section_map:
                continue
            section = section_map[section_id]
            detail_prompt = (
                "你已经理解整篇素材。现在只针对当前 section 返回局部编辑操作。"
                "严禁返回完整清洗后正文。"
                "允许的 operation.type 只有 delete_span、replace_span、insert_text、split_paragraph。"
                "replace_span 只能用于短替换，replacement 长度不得超过 120 字。"
                "删除目标包括：人物小传、人物介绍、剧情大纲、故事梗概、项目包装信息、导流信息、网址、作者对读者的话、乱码。"
                "必须保留章节号、集号、幕号、节号、分镜号，例如第1集、第3章、第2幕、第4节、分镜5。不要删除这些编号标题。"
                "如果当前 section 里存在明确正文起点（例如单独一行的1、01、001、1.、1、，或第1集、第一章、分镜1），必须返回一条 delete_span，从 start=0 删除到正文起点之前。"
                "正文起点之前的导语、前情、包装文案、人物小传、大纲都必须整体删除，不能只删一部分。"
                "需要补换行时，优先在句号、问号、叹号、引号前后补正文换行。"
                "如果需要大面积改写，只返回 issue 并 manual_review=true。"
                "返回 JSON 结构："
                '{"issues":[{"section_id":"...","category":"...","reason":"...","manual_review":false}],'
                '"deletion_summaries":[{"section_id":"...","category":"...","summary":"..."}],'
                '"operations":[{"type":"delete_span","section_id":"...","start":0,"end":10,"reason":"..."}]}'
            )
            detail_payload = {
                "section_id": section.section_id,
                "kind": section.kind,
                "locator": section.locator,
                "text": section.text,
                "categories": target.get("categories") or [],
                "reason": target.get("reason") or "",
            }
            detail_response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "只输出 JSON，不要输出解释、Markdown 或代码围栏。严禁输出完整清洗后正文。"},
                    {"role": "user", "content": [{"type": "text", "text": detail_prompt}]},
                    {"role": "user", "content": [{"type": "text", "text": json.dumps(detail_payload, ensure_ascii=False)}]},
                ],
                max_tokens=1500,
                stream=False,
                temperature=0,
            )
            detail_data = parse_json_object(extract_response_text(detail_response.model_dump()))
            operations.extend(detail_data.get("operations") or [])
            issues.extend(detail_data.get("issues") or [])
            deletion_summaries.extend(detail_data.get("deletion_summaries") or [])
        return {
            "issues": issues,
            "deletion_summaries": deletion_summaries,
            "operations": operations,
            "target_sections": targets,
        }


def extract_response_text(payload: dict[str, Any]) -> str:
    if isinstance(payload.get("output_text"), str) and payload["output_text"].strip():
        return payload["output_text"]
    output = payload.get("output")
    if isinstance(output, list):
        parts: list[str] = []
        for item in output:
            if not isinstance(item, dict):
                continue
            for content in item.get("content") or []:
                if isinstance(content, dict) and content.get("type") in {"output_text", "text"} and isinstance(content.get("text"), str):
                    parts.append(content["text"])
        if parts:
            return "\n".join(parts)
    choices = payload.get("choices") or []
    if not choices:
        raise CleaningError(f"model returned no text content: {json.dumps(payload, ensure_ascii=False)[:1000]}")
    message = choices[0].get("message") or {}
    content = message.get("content")
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict) and item.get("type") == "text" and isinstance(item.get("text"), str):
                parts.append(item["text"])
        if parts:
            return "\n".join(parts)
    raise CleaningError(f"model response did not contain text content: {json.dumps(payload, ensure_ascii=False)[:1000]}")


def parse_json_object(text: str) -> dict[str, Any]:
    stripped = text.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:json)?\s*|\s*```$", "", stripped, flags=re.IGNORECASE | re.DOTALL).strip()
    object_start = stripped.find("{")
    object_end = stripped.rfind("}")
    if object_start != -1 and object_end != -1 and object_end > object_start:
        stripped = stripped[object_start : object_end + 1]
    try:
        data = json.loads(stripped)
    except json.JSONDecodeError:
        data = repair_partial_json(stripped)
        if data is None:
            raise CleaningError(f"model returned invalid JSON; preview={text[:500]!r}")
    if not isinstance(data, dict):
        raise CleaningError("model output must be a JSON object")
    forbidden = {"cleaned_text", "rewritten_document", "final_text", "full_text"}
    overlap = forbidden.intersection(data.keys())
    if overlap:
        raise CleaningError(f"model returned forbidden full-text fields: {sorted(overlap)}")
    return data


def repair_partial_json(text: str) -> dict[str, Any] | None:
    issues: list[Any] = []
    deletions: list[Any] = []
    operations: list[dict[str, Any]] = []
    issues_match = re.search(r'"issues"\s*:\s*(\[.*?\])\s*,\s*"deletion_summaries"', text, flags=re.DOTALL)
    if issues_match:
        try:
            issues = json.loads(issues_match.group(1))
        except json.JSONDecodeError:
            issues = []
    deletions_match = re.search(r'"deletion_summaries"\s*:\s*(\[.*?\])\s*,\s*"operations"', text, flags=re.DOTALL)
    if deletions_match:
        deletions_text = deletions_match.group(1)
        try:
            deletions = json.loads(deletions_text)
        except json.JSONDecodeError:
            for snippet in re.findall(r'\{[^{}]*\}', deletions_text):
                try:
                    obj = json.loads(snippet)
                except json.JSONDecodeError:
                    continue
                deletions.append(obj)
    operations_match = re.search(r'"operations"\s*:\s*(\[.*)', text, flags=re.DOTALL)
    if operations_match:
        operations_text = operations_match.group(1)
        for snippet in re.findall(r'\{[^{}]*\}', operations_text):
            try:
                obj = json.loads(snippet)
            except json.JSONDecodeError:
                continue
            if isinstance(obj, dict):
                operations.append(obj)
    if not issues and not deletions and not operations:
        return None
    return {"issues": issues, "deletion_summaries": deletions, "operations": operations}


def preprocess_section_text(text: str) -> tuple[str, list[str]]:
    notes: list[str] = []
    current = normalize_text(text)
    if current != text:
        notes.append("normalized_line_endings")
    return current, notes


def chunk_text(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    chunks: list[str] = []
    current = ""
    for paragraph in re.split(r"\n{2,}", text):
        piece = paragraph.strip()
        if not piece:
            continue
        candidate = f"{current}\n\n{piece}" if current else piece
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current)
        if len(piece) <= max_chars:
            current = piece
            continue
        start = 0
        while start < len(piece):
            end = min(start + max_chars, len(piece))
            chunks.append(piece[start:end])
            start = end
        current = ""
    if current:
        chunks.append(current)
    return chunks or [text[:max_chars]]


def parse_column_reference(value: str) -> int:
    stripped = value.strip()
    if not stripped:
        raise CleaningError("column reference must not be empty")
    if stripped.isdigit():
        column_index = int(stripped)
        if column_index <= 0:
            raise CleaningError("column index must be greater than 0")
        return column_index
    column_index = 0
    for char in stripped.upper():
        if char < "A" or char > "Z":
            raise CleaningError(f"invalid column reference: {value}")
        column_index = column_index * 26 + (ord(char) - ord("A") + 1)
    return column_index


def resolve_candidate_columns(sheet, requested_column: str | None, requested_header: str | None) -> tuple[list[int], dict[int, str]]:
    header_cells = next(sheet.iter_rows(min_row=1, max_row=1, values_only=False), [])
    header_map = {idx: ("" if cell.value is None else str(cell.value)) for idx, cell in enumerate(header_cells, start=1)}
    if requested_column:
        column_index = parse_column_reference(requested_column)
        if column_index > (sheet.max_column or 1):
            raise CleaningError(f"sheet {sheet.title} does not have column {requested_column}")
        return [column_index], header_map
    if requested_header:
        lowered = requested_header.strip().lower()
        matches = [idx for idx, value in header_map.items() if value.strip().lower() == lowered]
        if not matches:
            raise CleaningError(f"sheet {sheet.title} does not have header {requested_header!r}")
        return matches, header_map
    candidate_columns: list[int] = []
    for idx, value in header_map.items():
        lowered = value.lower()
        if any(token in lowered for token in ["response", "text", "content", "正文", "内容", "文案", "素材"]):
            candidate_columns.append(idx)
    if candidate_columns:
        return candidate_columns, header_map
    max_col = sheet.max_column or 1
    if max_col == 1:
        return [1], header_map
    string_counts: list[tuple[int, int]] = []
    for col in range(1, max_col + 1):
        count = 0
        for row in range(2, min(sheet.max_row, 20) + 1):
            value = sheet.cell(row=row, column=col).value
            if isinstance(value, str) and value.strip():
                count += 1
        string_counts.append((col, count))
    best = max(string_counts, key=lambda item: item[1])[0]
    return [best], header_map


def first_body_start_section_id(sections: list[Section]) -> str | None:
    for section in sections:
        if section.text.strip() and BODY_START_MARK_RE.match(section.text.strip()):
            return section.section_id
    return None


def load_material(path: Path, *, sheet_name: str | None = None, column: str | None = None, header_name: str | None = None) -> MaterialDocument:
    suffix = path.suffix.lower()
    if suffix == ".doc":
        raise CleaningError("unsupported file type: .doc (please convert to .docx first)")
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8")
        processed, notes = preprocess_section_text(text)
        sections = [Section(section_id="txt:0", text=processed, original_text=text, preprocess_notes=notes, kind="text_block", locator={"paragraph_index": 0})]
        return MaterialDocument(source_path=path, source_type="txt", sections=sections, selection={"sheet_name": sheet_name, "column": column, "header_name": header_name})
    if suffix == ".docx":
        doc = Document(str(path))
        sections: list[Section] = []
        for index, paragraph in enumerate(doc.paragraphs):
            processed, notes = preprocess_section_text(paragraph.text)
            sections.append(
                Section(
                    section_id=f"docx:{index}",
                    text=processed,
                    original_text=paragraph.text,
                    preprocess_notes=notes,
                    kind="paragraph",
                    locator={"paragraph_index": index},
                )
            )
        return MaterialDocument(
            source_path=path,
            source_type="docx",
            sections=sections,
            selection={"sheet_name": sheet_name, "column": column, "header_name": header_name},
            body_start_section_id=first_body_start_section_id(sections),
        )
    if suffix == ".xlsx":
        workbook = load_workbook(path)
        sections: list[Section] = []
        editable_columns: dict[str, list[int]] = {}
        header_map: dict[str, dict[int, str]] = {}
        worksheets = workbook.worksheets
        if sheet_name is not None:
            if sheet_name not in workbook.sheetnames:
                raise CleaningError(f"sheet not found: {sheet_name}")
            worksheets = [workbook[sheet_name]]
        for sheet in worksheets:
            candidate_columns, sheet_headers = resolve_candidate_columns(sheet, column, header_name)
            editable_columns[sheet.title] = candidate_columns
            header_map[sheet.title] = sheet_headers
            for row in range(2, sheet.max_row + 1):
                for col in candidate_columns:
                    value = sheet.cell(row=row, column=col).value
                    if isinstance(value, str):
                        processed, notes = preprocess_section_text(value)
                        chunks = chunk_text(processed, DEFAULT_MAX_SECTION_CHARS)
                        for chunk_index, chunk in enumerate(chunks):
                            sections.append(
                                Section(
                                    section_id=f"xlsx:{sheet.title}:{row}:{col}:{chunk_index}",
                                    text=chunk,
                                    original_text=value,
                                    preprocess_notes=notes,
                                    kind="cell_chunk",
                                    locator={
                                        "sheet_name": sheet.title,
                                        "row_index": row,
                                        "column_index": col,
                                        "column_header": sheet_headers.get(col, ""),
                                        "chunk_index": chunk_index,
                                        "chunk_count": len(chunks),
                                    },
                                )
                            )
        return MaterialDocument(
            source_path=path,
            source_type="xlsx",
            sections=sections,
            workbook=workbook,
            sheet_order=[sheet.title for sheet in worksheets],
            editable_columns=editable_columns,
            header_map=header_map,
            selection={"sheet_name": sheet_name, "column": column, "header_name": header_name},
        )
    raise CleaningError(f"unsupported file type: {suffix}")


def add_heuristic_operations(document: MaterialDocument, model_output: dict[str, Any]) -> dict[str, Any]:
    operations = [op for op in list(model_output.get("operations") or []) if not deletes_chapter_marker(document, op)]
    issues = list(model_output.get("issues") or [])
    deletion_summaries = list(model_output.get("deletion_summaries") or [])
    existing_keys = {
        (op.get("type"), op.get("section_id"), op.get("start"), op.get("end"), op.get("offset"))
        for op in operations
        if isinstance(op, dict)
    }
    for section in document.sections:
        text = section.text
        if not text:
            continue
        chapter_ranges = chapter_marker_ranges(text)
        metadata_end = leading_metadata_end(text)
        if metadata_end > 0:
            key = ("delete_span", section.section_id, 0, metadata_end, None)
            if key not in existing_keys:
                operations.append({"type": "delete_span", "section_id": section.section_id, "start": 0, "end": metadata_end, "reason": "remove leading metadata block"})
                deletion_summaries.append({"section_id": section.section_id, "category": "leading_metadata", "summary": text[: min(metadata_end, 40)]})
                existing_keys.add(key)
        lines = text.splitlines()
        for line in lines[:8]:
            stripped = line.strip()
            if not stripped or CHAPTER_MARK_RE.match(stripped):
                continue
            if TITLEISH_RE.match(stripped):
                start = text.find(line)
                end = start + len(line)
                if overlaps_ranges(start, end, chapter_ranges):
                    continue
                key = ("delete_span", section.section_id, start, end, None)
                if key not in existing_keys:
                    operations.append({"type": "delete_span", "section_id": section.section_id, "start": start, "end": end, "reason": "remove leading metadata"})
                    deletion_summaries.append({"section_id": section.section_id, "category": "leading_metadata", "summary": stripped[:40]})
                    existing_keys.add(key)
        for match in URL_RE.finditer(text):
            key = ("delete_span", section.section_id, match.start(), match.end(), None)
            if key not in existing_keys:
                operations.append({"type": "delete_span", "section_id": section.section_id, "start": match.start(), "end": match.end(), "reason": "remove url"})
                deletion_summaries.append({"section_id": section.section_id, "category": "url", "summary": match.group(0)[:40]})
                existing_keys.add(key)
        for match in GARBLED_RE.finditer(text):
            key = ("delete_span", section.section_id, match.start(), match.end(), None)
            if key not in existing_keys:
                operations.append({"type": "delete_span", "section_id": section.section_id, "start": match.start(), "end": match.end(), "reason": "remove garbled characters"})
                deletion_summaries.append({"section_id": section.section_id, "category": "garbled", "summary": match.group(0)[:20]})
                existing_keys.add(key)
        for pattern in TRAILING_PATTERNS:
            for match in pattern.finditer(text):
                if overlaps_ranges(match.start(), match.end(), chapter_ranges):
                    continue
                key = ("delete_span", section.section_id, match.start(), match.end(), None)
                if key not in existing_keys:
                    operations.append({"type": "delete_span", "section_id": section.section_id, "start": match.start(), "end": match.end(), "reason": "remove trailing metadata"})
                    deletion_summaries.append({"section_id": section.section_id, "category": "trailing_metadata", "summary": match.group(0)[:40]})
                    existing_keys.add(key)
        for word in BUTTON_WORDS:
            start = text.find(word)
            while start != -1:
                end = start + len(word)
                if not overlaps_ranges(start, end, chapter_ranges):
                    key = ("delete_span", section.section_id, start, end, None)
                    if key not in existing_keys:
                        operations.append({"type": "delete_span", "section_id": section.section_id, "start": start, "end": end, "reason": "remove button label"})
                        deletion_summaries.append({"section_id": section.section_id, "category": "button_label", "summary": word})
                        existing_keys.add(key)
                start = text.find(word, end)
        for offset in newline_offsets(text):
            key = ("insert_text", section.section_id, None, None, offset)
            if key not in existing_keys:
                operations.append({"type": "insert_text", "section_id": section.section_id, "offset": offset, "text": "\n", "reason": "insert paragraph break"})
                existing_keys.add(key)
        if any(token in text for token in ["作者有话说", "读者朋友", "宝子们", "求收藏", "求关注"]):
            issues.append({"section_id": section.section_id, "category": "author_aside", "reason": "contains possible author-to-reader aside", "manual_review": True})
    model_output["operations"] = operations
    model_output["issues"] = issues
    model_output["deletion_summaries"] = deletion_summaries
    return model_output


def chapter_marker_ranges(text: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for match in CHAPTER_MARK_RE.finditer(text):
        ranges.append((match.start(), match.end()))
    return ranges


def overlaps_ranges(start: int, end: int, ranges: list[tuple[int, int]]) -> bool:
    for range_start, range_end in ranges:
        if not (end <= range_start or start >= range_end):
            return True
    return False


def leading_metadata_end(text: str) -> int:
    if not text.strip():
        return 0
    anchor_end = body_start_anchor_end(text)
    if anchor_end > 0:
        return anchor_end
    lines = text.splitlines(keepends=True)
    offset = 0
    found_metadata = False
    for line in lines[:120]:
        stripped = line.strip()
        if not stripped:
            offset += len(line)
            continue
        if CHAPTER_MARK_RE.match(stripped) or looks_like_story_start(stripped):
            return offset if found_metadata else 0
        if found_metadata:
            offset += len(line)
            continue
        if METADATA_BLOCK_START_RE.match(stripped) or TITLEISH_RE.match(stripped):
            found_metadata = True
            offset += len(line)
            continue
        if looks_like_platform_noise(stripped):
            found_metadata = True
            offset += len(line)
            continue
        break
    return offset if found_metadata else 0


def body_start_anchor_end(text: str) -> int:
    if not text.strip():
        return 0
    lines = text.splitlines(keepends=True)
    offset = 0
    first_story_offset: int | None = None
    for line in lines[:220]:
        stripped = line.strip()
        if not stripped:
            offset += len(line)
            continue
        if BODY_START_MARK_RE.match(stripped):
            return offset
        if first_story_offset is None and looks_like_story_start(stripped):
            first_story_offset = offset
        offset += len(line)
    return first_story_offset or 0


def newline_offsets(text: str) -> list[int]:
    if not text:
        return []
    offsets: list[int] = []
    for match in re.finditer(r"([。！？；])([“\"‘'（(第分A-Za-z\u4e00-\u9fff])", text):
        next_char = match.group(2)
        if next_char and next_char not in "，。；：！？、)）】]}》’'\n":
            offsets.append(match.start(2))
    for match in re.finditer(r"(】|）|》)([第分A-Za-z\u4e00-\u9fff])", text):
        offsets.append(match.start(2))
    for match in re.finditer(r"(第\s*[0-9一二三四五六七八九十百千两零]+\s*[集章节幕回卷]|分镜\s*[0-9一二三四五六七八九十百千两零]+)", text):
        if match.start() > 0:
            offsets.append(match.start())
    return [offset for offset in sorted(set(offsets)) if 0 < offset < len(text)]


def looks_like_platform_noise(line: str) -> bool:
    return line in {"无障碍", "查看详情", "知乎盐选"} or line.startswith("知乎盐选")


def looks_like_story_start(line: str) -> bool:
    if CHAPTER_MARK_RE.match(line) or BODY_START_MARK_RE.match(line):
        return True
    if line in {"正文", "正文开始"}:
        return True
    return bool(re.match(r"^(我|他|她|他们|她们|那天|那年|后来|当时|此时|这一年|一天|两个月前|1[.、]?|01|001)", line))


def deletes_chapter_marker(document: MaterialDocument, op: dict[str, Any]) -> bool:
    if not isinstance(op, dict):
        return False
    if op.get("type") not in {"delete_span", "replace_span"}:
        return False
    section_id = op.get("section_id")
    start = op.get("start")
    end = op.get("end")
    if not isinstance(section_id, str) or not isinstance(start, int) or not isinstance(end, int):
        return False
    section = next((item for item in document.sections if item.section_id == section_id), None)
    if section is None:
        return False
    snippet = section.text[start:end].strip()
    return bool(snippet and CHAPTER_MARK_RE.match(snippet))


def validate_operations(document: MaterialDocument, model_output: dict[str, Any], max_replacement_length: int, max_edit_ratio: float) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    section_map = {section.section_id: section for section in document.sections}
    accepted: list[dict[str, Any]] = []
    rejected: list[dict[str, Any]] = []
    spans_by_section: dict[str, list[tuple[int, int]]] = {}
    edit_load: dict[str, int] = {}
    for raw_op in model_output.get("operations") or []:
        if not isinstance(raw_op, dict):
            rejected.append({"operation": raw_op, "reason": "operation must be an object"})
            continue
        op_type = raw_op.get("type")
        section_id = raw_op.get("section_id")
        if not isinstance(section_id, str) or section_id not in section_map:
            rejected.append({"operation": raw_op, "reason": "unknown section_id"})
            continue
        section = section_map[section_id]
        text_len = len(section.text)
        if op_type == "delete_span" or op_type == "replace_span":
            start = raw_op.get("start")
            end = raw_op.get("end")
            if not isinstance(start, int) or not isinstance(end, int) or start < 0 or end < start or end > text_len:
                rejected.append({"operation": raw_op, "reason": "invalid span"})
                continue
            for existing_start, existing_end in spans_by_section.setdefault(section_id, []):
                if not (end <= existing_start or start >= existing_end):
                    rejected.append({"operation": raw_op, "reason": "overlapping span"})
                    break
            else:
                if op_type == "replace_span":
                    replacement = raw_op.get("replacement")
                    if not isinstance(replacement, str):
                        rejected.append({"operation": raw_op, "reason": "replacement must be string"})
                        continue
                    if len(replacement) > max_replacement_length:
                        rejected.append({"operation": raw_op, "reason": "replacement too long"})
                        continue
                    if len(replacement) > 20:
                        rejected.append({"operation": raw_op, "reason": "replacement too long for no-rewrite policy"})
                        continue
                changed = end - start
                if not (start == 0 and end == text_len):
                    edit_load[section_id] = edit_load.get(section_id, 0) + changed
                    if text_len and (edit_load[section_id] / text_len) > max_edit_ratio:
                        rejected.append({"operation": raw_op, "reason": "edit ratio too high"})
                        edit_load[section_id] -= changed
                        continue
                spans_by_section[section_id].append((start, end))
                accepted.append(raw_op)
                continue
            continue
        if op_type == "insert_text" or op_type == "split_paragraph":
            offset = raw_op.get("offset")
            if not isinstance(offset, int) or offset < 0 or offset > text_len:
                rejected.append({"operation": raw_op, "reason": "invalid offset"})
                continue
            if op_type == "insert_text":
                insert_text = raw_op.get("text")
                if not isinstance(insert_text, str) or len(insert_text) > max_replacement_length:
                    rejected.append({"operation": raw_op, "reason": "insert text invalid or too long"})
                    continue
                if insert_text != "\n":
                    rejected.append({"operation": raw_op, "reason": "insert_text may only insert newline"})
                    continue
            accepted.append(raw_op)
            continue
        rejected.append({"operation": raw_op, "reason": "unsupported operation type"})
    return accepted, rejected


def apply_operations(document: MaterialDocument, operations: list[dict[str, Any]]) -> dict[str, str]:
    by_section: dict[str, list[dict[str, Any]]] = {}
    for op in operations:
        by_section.setdefault(op["section_id"], []).append(op)
    result: dict[str, str] = {}
    body_start_seen = document.body_start_section_id is None
    for section in document.sections:
        text = section.text
        if document.source_type == "docx" and not body_start_seen:
            if section.section_id == document.body_start_section_id:
                body_start_seen = True
            else:
                result[section.section_id] = ""
                continue
        anchor_trim = 0
        if document.source_type == "xlsx":
            anchor_trim = body_start_anchor_end(text)
            if anchor_trim > 0:
                text = text[anchor_trim:]
        section_ops = by_section.get(section.section_id, [])
        section_ops.sort(key=sort_key_for_apply, reverse=True)
        for op in section_ops:
            op_type = op["type"]
            if op_type == "delete_span":
                start = max(op["start"] - anchor_trim, 0) if anchor_trim else op["start"]
                end = max(op["end"] - anchor_trim, 0) if anchor_trim else op["end"]
                if end > len(text) or start > len(text) or end < start:
                    continue
                text = text[:start] + text[end:]
            elif op_type == "replace_span":
                start = max(op["start"] - anchor_trim, 0) if anchor_trim else op["start"]
                end = max(op["end"] - anchor_trim, 0) if anchor_trim else op["end"]
                if end > len(text) or start > len(text) or end < start:
                    continue
                text = text[:start] + op["replacement"] + text[end:]
            elif op_type == "insert_text":
                offset = max(op["offset"] - anchor_trim, 0) if anchor_trim else op["offset"]
                if offset > len(text):
                    continue
                text = text[:offset] + op["text"] + text[offset:]
            elif op_type == "split_paragraph":
                offset = max(op["offset"] - anchor_trim, 0) if anchor_trim else op["offset"]
                if offset > len(text):
                    continue
                text = text[:offset] + "\n" + text[offset:]
        result[section.section_id] = cleanup_text(text)
    return result


def sort_key_for_apply(op: dict[str, Any]) -> tuple[int, int]:
    if "start" in op:
        return int(op["start"]), 1
    return int(op.get("offset", 0)), 0


def normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def cleanup_text(text: str) -> str:
    text = normalize_text(text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def write_outputs(document: MaterialDocument, cleaned_sections: dict[str, str], report: dict[str, Any], raw_model_output: dict[str, Any]) -> dict[str, str]:
    source = document.source_path
    cleaned_path = source.with_name(f"{source.stem}.cleaned{source.suffix}")
    report_path = source.with_name(f"{source.stem}.cleaning-report.json")
    raw_output_path = source.with_name(f"{source.stem}.raw-model-output.json")
    if document.source_type == "txt":
        cleaned_path.write_text(cleaned_sections.get("txt:0", ""), encoding="utf-8")
    elif document.source_type == "docx":
        source_doc = Document(str(source))
        for index, paragraph in enumerate(source_doc.paragraphs):
            paragraph.text = cleaned_sections.get(f"docx:{index}", paragraph.text)
        source_doc.save(cleaned_path)
    elif document.source_type == "xlsx":
        workbook = copy.deepcopy(document.workbook)
        grouped: dict[tuple[str, int, int], list[Section]] = {}
        for section in document.sections:
            locator = section.locator
            key = (locator["sheet_name"], locator["row_index"], locator["column_index"])
            grouped.setdefault(key, []).append(section)
        for (sheet_name, row_index, column_index), sections in grouped.items():
            ordered = sorted(sections, key=lambda item: item.locator.get("chunk_index", 0))
            merged = "\n\n".join(cleaned_sections.get(section.section_id, section.text) for section in ordered if cleaned_sections.get(section.section_id, section.text))
            workbook[sheet_name].cell(row=row_index, column=column_index).value = merged
        workbook.save(cleaned_path)
    else:
        raise CleaningError(f"unsupported output type: {document.source_type}")
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    raw_output_path.write_text(json.dumps(raw_model_output, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "cleaned": str(cleaned_path),
        "report": str(report_path),
        "raw_model_output": str(raw_output_path),
    }


def build_report(document: MaterialDocument, accepted: list[dict[str, Any]], rejected: list[dict[str, Any]], model_output: dict[str, Any], cleaned_sections: dict[str, str]) -> dict[str, Any]:
    section_map = {section.section_id: section for section in document.sections}
    changed_sections = []
    issue_counts = Counter(issue.get("category") for issue in (model_output.get("issues") or []) if isinstance(issue, dict) and issue.get("category"))
    deletion_counts = Counter(item.get("category") for item in (model_output.get("deletion_summaries") or []) if isinstance(item, dict) and item.get("category"))
    manual_review_sections = sorted(
        {
            issue.get("section_id")
            for issue in (model_output.get("issues") or [])
            if isinstance(issue, dict) and issue.get("manual_review") and isinstance(issue.get("section_id"), str)
        }
    )
    for section_id, cleaned in cleaned_sections.items():
        original = section_map[section_id].text
        if cleaned != cleanup_text(original):
            changed_sections.append(
                {
                    "section_id": section_id,
                    "locator": section_map[section_id].locator,
                    "original_length": len(original),
                    "cleaned_length": len(cleaned),
                }
            )
    return {
        "source_path": str(document.source_path),
        "source_type": document.source_type,
        "section_count": len(document.sections),
        "changed_section_count": len(changed_sections),
        "unchanged_section_count": len(document.sections) - len(changed_sections),
        "manual_review_count": len(manual_review_sections),
        "manual_review_sections": manual_review_sections,
        "selection": document.selection or {},
        "editable_columns": document.editable_columns or {},
        "header_map": document.header_map or {},
        "accepted_operation_count": len(accepted),
        "rejected_operation_count": len(rejected),
        "issue_category_counts": dict(issue_counts),
        "deletion_category_counts": dict(deletion_counts),
        "issues": model_output.get("issues") or [],
        "deletion_summaries": model_output.get("deletion_summaries") or [],
        "accepted_operations": accepted,
        "rejected_operations": rejected,
        "changed_sections": changed_sections,
    }


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Clean novel/script materials without letting the model emit full cleaned text.")
    parser.add_argument("input", help="Path to txt, docx, or xlsx material")
    parser.add_argument("--api-key", default=os.getenv("MATERIAL_CLEANER_API_KEY") or os.getenv("AI_TEXT_API_KEY"))
    parser.add_argument("--endpoint", default=os.getenv("MATERIAL_CLEANER_ENDPOINT", DEFAULT_ENDPOINT))
    parser.add_argument("--api-version", default=os.getenv("MATERIAL_CLEANER_API_VERSION", DEFAULT_API_VERSION))
    parser.add_argument("--model", default=os.getenv("MATERIAL_CLEANER_MODEL", DEFAULT_MODEL))
    parser.add_argument("--logid", default=os.getenv("MATERIAL_CLEANER_LOGID"))
    parser.add_argument("--sheet")
    parser.add_argument("--column")
    parser.add_argument("--header-name")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--report-only", action="store_true")
    parser.add_argument("--max-replacement-length", type=int, default=DEFAULT_MAX_REPLACEMENT_LENGTH)
    parser.add_argument("--max-edit-ratio", type=float, default=DEFAULT_MAX_EDIT_RATIO)
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    source_path = Path(args.input).expanduser().resolve()
    if not source_path.exists():
        print(f"input not found: {source_path}", file=sys.stderr)
        return 1
    if not args.api_key:
        print("missing API key: pass --api-key or set MATERIAL_CLEANER_API_KEY/AI_TEXT_API_KEY", file=sys.stderr)
        return 1
    try:
        document = load_material(source_path, sheet_name=args.sheet, column=args.column, header_name=args.header_name)
        client = AzureCleaningClient(
            api_key=args.api_key,
            endpoint=args.endpoint,
            api_version=args.api_version,
            model=args.model,
            logid=args.logid,
        )
        raw_model_output = client.analyze_document(document)
        enriched_output = add_heuristic_operations(document, raw_model_output)
        accepted, rejected = validate_operations(document, enriched_output, args.max_replacement_length, args.max_edit_ratio)
        cleaned_sections = apply_operations(document, accepted)
        report = build_report(document, accepted, rejected, enriched_output, cleaned_sections)
        if args.report_only or args.dry_run:
            print(json.dumps({"report": report, "raw_model_output": raw_model_output}, ensure_ascii=False, indent=2))
            return 0
        outputs = write_outputs(document, cleaned_sections, report, raw_model_output)
        print(json.dumps(outputs, ensure_ascii=False, indent=2))
        return 0
    except Exception as exc:
        print(f"cleaning failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
